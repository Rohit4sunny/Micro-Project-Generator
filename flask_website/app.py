from flask import Flask, render_template, request, send_file
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from dotenv import load_dotenv
import google.generativeai as genai
from requests import get
from bs4 import BeautifulSoup

load_dotenv()  # Load environment variables from .env file

# Initialize Gemini API
genai.configure(api_key=os.environ["GEMINI_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-flash')

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_report():
    title = request.form['title']
    document = Document()

    # Title
    title_heading = document.add_heading(title, level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Content from Gemini API
    content = fetch_content(title)
    if content:
        paragraphs = process_content(document, content)
        images = fetch_images(title)

        # Insert images at suitable places within the paragraphs
        insert_images(document, paragraphs, images)
    else:
        document.add_paragraph("No content available for this topic.")

    # Save the document
    temp_dir = 'tmp'
    os.makedirs(temp_dir, exist_ok=True)
    filepath = os.path.join(temp_dir, f"{title}.docx")
    document.save(filepath)
    return send_file(filepath, as_attachment=True)

def fetch_content(title):
    try:
        response = model.generate_content(f"Generate a detailed and professional Micro Project Report on {title} with proper structure, suitable for engineering students. Include sections such as Introduction, Working Principle, Methodology, Classification, Applications, Results, Conclusion, and References.")
        
        # Debugging information
        print(f"Response: {response}")

        if response and response.text:
            return response.text
        else:
            return "No content available for this topic."
    except Exception as e:
        print(f"Error fetching content: {e}")
    return "No content available for this topic."

def process_content(document, content):
    paragraphs = []
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith("## "):
            heading = document.add_heading(line[3:], level=1)
            for run in heading.runs:
                run.font.size = Pt(18)
                run.font.name = 'Arial'
        elif line.startswith("### "):
            heading = document.add_heading(line[4:], level=2)
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.name = 'Arial'
        elif line.startswith("* "):
            heading = document.add_heading(line[2:], level=3)
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.name = 'Arial'
        else:
            p = document.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
            paragraphs.append(p)
    return paragraphs

def fetch_images(title):
    try:
        response = get(f"https://www.google.com/search?tbm=isch&q={title}")
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            img_tags = soup.find_all('img')
            img_urls = []
            for img in img_tags:
                img_url = img.get('src')
                if img_url and img_url.startswith('http'):
                    img_urls.append(img_url)
            return img_urls[:5]  # Return the first 5 valid image URLs
    except Exception as e:
        print(f"Error fetching images: {e}")
    return []

def insert_images(document, paragraphs, images):
    for i, paragraph in enumerate(paragraphs):
        if i % 5 == 0 and i // 5 < len(images):
            img_url = images[i // 5]
            response = get(img_url)
            if response.status_code == 200:
                image_path = f'tmp/image_{i // 5}.jpg'
                with open(image_path, 'wb') as img_file:
                    img_file.write(response.content)
                run = paragraph.add_run()
                run.add_break()
                run.add_picture(image_path, width=Inches(4.0))
                os.remove(image_path)  # Clean up image after adding

if __name__ == '__main__':
    app.run(debug=True)
